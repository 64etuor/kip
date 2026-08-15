from __future__ import annotations

import zipfile
from collections.abc import Mapping
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from kip.adapters.parsers import docx as docx_module
from kip.adapters.parsers.docx import DocxParser
from kip.errors import ParserError, ValidationError


def _replace_zip_members(path: Path, replacements: Mapping[str, bytes]) -> None:
    with zipfile.ZipFile(path) as archive:
        members = {name: archive.read(name) for name in archive.namelist()}
    members.update(replacements)
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, content in members.items():
            archive.writestr(name, content)


def _parse(path: Path, **overrides: object):
    kwargs: dict[str, object] = {
        "artifact_id": "art_docx",
        "document_id": "doc_docx",
        "acl_scopes": ["workspace:default"],
    }
    kwargs.update(overrides)
    return DocxParser().parse(path, **kwargs)  # type: ignore[arg-type]


def test_docx_parser_captures_heading_levels(tmp_path: Path) -> None:
    # Given a document mixing title-level and section-level headings with body text.
    path = tmp_path / "headings.docx"
    document = Document()
    document.add_heading("제목1", level=1)
    document.add_paragraph("본문 단락")
    document.add_heading("소제목", level=2)
    document.save(path)

    # When the document is parsed.
    _extraction, units = _parse(path)

    # Then the heading text, level, and body paragraph position are recorded.
    assert len(units) == 1
    assert units[0].unit_type == "docx_paragraph"
    assert units[0].metadata["headings"] == [
        {"text": "제목1", "level": 1, "paragraph": 1},
        {"text": "소제목", "level": 2, "paragraph": 3},
    ]
    assert units[0].locator.data == {"start_paragraph": 1, "end_paragraph": 3}


def test_docx_parser_prefixes_list_paragraphs(tmp_path: Path) -> None:
    # Given a paragraph carrying explicit list numbering properties.
    path = tmp_path / "list.docx"
    document = Document()
    paragraph = document.add_paragraph("목록 항목")
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.makeelement(qn("w:numPr"), {})
    ilvl = num_pr.makeelement(qn("w:ilvl"), {qn("w:val"): "0"})
    num_id = num_pr.makeelement(qn("w:numId"), {qn("w:val"): "1"})
    num_pr.append(ilvl)
    num_pr.append(num_id)
    ppr.append(num_pr)
    document.save(path)

    # When the document is parsed.
    _extraction, units = _parse(path)

    # Then the paragraph text is prefixed with a simple list marker.
    assert units[0].body == "- 목록 항목"


def test_docx_parser_renders_table_with_gridspan_and_vmerge_without_duplication(
    tmp_path: Path,
) -> None:
    # Given a table with a vertical merge and a horizontal merge, where the
    # vMerge continuation cell has been tampered with to carry duplicate text.
    path = tmp_path / "table.docx"
    document = Document()
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "일정"
    table.cell(1, 0).text = "계획"
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(2, 0).text = "실적"
    table.cell(2, 1).text = "완료"
    table.cell(2, 0).merge(table.cell(2, 1))
    document.save(path)
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")

    # Simulate a non-Word producer that leaves duplicate text on the vMerge
    # continuation cell instead of an empty paragraph.
    reopened = Document(path)
    reopened_table = reopened.tables[0]
    continuation_tc = reopened_table.rows[1]._tr.findall(qn("w:tc"))[0]
    paragraph_elem = continuation_tc.find(qn("w:p"))
    run_elem = paragraph_elem.makeelement(qn("w:r"), {})
    text_elem = run_elem.makeelement(qn("w:t"), {})
    text_elem.text = "DUPLICATE"
    run_elem.append(text_elem)
    paragraph_elem.append(run_elem)
    reopened.save(path)
    del document_xml

    # When the tampered table is parsed.
    _extraction, units = _parse(path)

    # Then the table renders as one tab-delimited row per source row, the
    # vMerge continuation cell stays empty regardless of tampered XML text,
    # and the horizontal merge pads the spanned column instead of repeating text.
    table_units = [unit for unit in units if unit.unit_type == "docx_table"]
    assert len(table_units) == 1
    assert table_units[0].locator.data == {"table_index": 0}
    assert table_units[0].body == "일정\n계획\t\n\t\n실적\n완료\t"
    assert "DUPLICATE" not in table_units[0].body
    assert table_units[0].metadata == {
        "row_count": 3,
        "col_count": 2,
        "nested_table_count": 0,
    }


def test_docx_parser_emits_nonempty_header_and_footer_units(tmp_path: Path) -> None:
    # Given a document with header and footer text.
    path = tmp_path / "header_footer.docx"
    document = Document()
    document.add_paragraph("본문")
    document.sections[0].header.add_paragraph("머리말 내용")
    document.sections[0].footer.add_paragraph("바닥글 내용")
    document.save(path)

    # When the document is parsed.
    _extraction, units = _parse(path)

    # Then one unit is emitted per non-empty header/footer part.
    header_footer_units = {
        unit.locator.data["part"]: unit for unit in units if unit.unit_type == "docx_header_footer"
    }
    assert set(header_footer_units) == {"word/header1.xml", "word/footer1.xml"}
    assert "머리말 내용" in header_footer_units["word/header1.xml"].body
    assert "바닥글 내용" in header_footer_units["word/footer1.xml"].body
    assert header_footer_units["word/header1.xml"].metadata["part_type"] == "header"
    assert header_footer_units["word/footer1.xml"].metadata["part_type"] == "footer"


def test_docx_parser_resolves_hyperlink_target_in_metadata(tmp_path: Path) -> None:
    # Given a document whose paragraph will carry a hyperlink relationship.
    path = tmp_path / "hyperlink.docx"
    document = Document()
    document.add_paragraph("LINKPLACEHOLDER")
    document.save(path)
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")
        rels_xml = archive.read("word/_rels/document.xml.rels")
    document_xml = document_xml.replace(
        b"<w:p><w:r><w:t>LINKPLACEHOLDER</w:t></w:r></w:p>",
        b'<w:p><w:hyperlink r:id="rIdLink"><w:r><w:t>\xeb\xb0\x94\xeb\xa1\x9c\xea\xb0\x80\xea\xb8\xb0</w:t></w:r></w:hyperlink></w:p>',
    )
    rels_xml = rels_xml.replace(
        b"</Relationships>",
        (
            b'<Relationship Id="rIdLink" '
            b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" '
            b'Target="https://example.com/kip" TargetMode="External"/></Relationships>'
        ),
    )
    _replace_zip_members(
        path, {"word/document.xml": document_xml, "word/_rels/document.xml.rels": rels_xml}
    )

    # When the document is parsed.
    _extraction, units = _parse(path)

    # Then the anchor text stays in the body and the resolved target is recorded.
    assert units[0].body == "바로가기"
    assert units[0].metadata["links"] == [
        {"text": "바로가기", "target": "https://example.com/kip"}
    ]


def test_docx_parser_extracts_textbox_once_and_excludes_it_from_body(tmp_path: Path) -> None:
    # Given a paragraph replaced by an mc:AlternateContent text box that Word
    # emits twice: once as DrawingML under mc:Choice, once as legacy VML under
    # mc:Fallback.
    path = tmp_path / "textbox.docx"
    document = Document()
    document.add_paragraph("TEXTBOXPLACEHOLDER")
    document.add_paragraph("일반 본문")
    document.save(path)
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")
    textbox_text = "텍스트상자내용"
    alternate_content = (
        "<w:p><w:r><mc:AlternateContent>"
        '<mc:Choice Requires="wps">'
        "<w:drawing><wp:anchor><wp:docPr id=\"1\" name=\"TextBox 1\"/>"
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        "<wps:wsp><wps:txbx><w:txbxContent>"
        f"<w:p><w:r><w:t>{textbox_text}</w:t></w:r></w:p>"
        "</w:txbxContent></wps:txbx></wps:wsp>"
        "</a:graphicData></a:graphic></wp:anchor></w:drawing>"
        "</mc:Choice>"
        "<mc:Fallback>"
        "<w:pict><v:shape><v:textbox><w:txbxContent>"
        f"<w:p><w:r><w:t>{textbox_text}</w:t></w:r></w:p>"
        "</w:txbxContent></v:textbox></v:shape></w:pict>"
        "</mc:Fallback>"
        "</mc:AlternateContent></w:r></w:p>"
    ).encode()
    document_xml = document_xml.replace(
        b"<w:p><w:r><w:t>TEXTBOXPLACEHOLDER</w:t></w:r></w:p>", alternate_content
    )
    _replace_zip_members(path, {"word/document.xml": document_xml})

    # When the document is parsed.
    _extraction, units = _parse(path)

    # Then exactly one text box unit is emitted (Choice/Fallback de-duplicated)
    # and its content never appears inside a body paragraph unit.
    textbox_units = [unit for unit in units if unit.unit_type == "docx_textbox"]
    assert len(textbox_units) == 1
    assert textbox_units[0].body == textbox_text
    assert textbox_units[0].locator.data == {"textbox_index": 0}
    paragraph_units = [unit for unit in units if unit.unit_type == "docx_paragraph"]
    assert all(textbox_text not in unit.body for unit in paragraph_units)
    assert any("일반 본문" in unit.body for unit in paragraph_units)


def test_docx_parser_counts_embedded_images(tmp_path: Path) -> None:
    # Given a document with one embedded PNG image.
    image_path = tmp_path / "figure.png"
    Image.new("RGB", (4, 4), color="white").save(image_path)
    path = tmp_path / "image.docx"
    document = Document()
    document.add_paragraph("본문")
    document.add_picture(str(image_path))
    document.save(path)

    # When the document is parsed.
    extraction, _units = _parse(path)

    # Then the image is counted with its relationship target and content type,
    # without reading or hashing image bytes.
    assert extraction.metadata["image_count"] == 1
    assert extraction.metadata["images"] == [
        {"target": "word/media/image1.png", "content_type": "image/png"}
    ]


def test_docx_parser_keeps_primary_content_when_header_part_is_malformed(tmp_path: Path) -> None:
    # Given a valid document whose header part has been corrupted.
    path = tmp_path / "malformed_header.docx"
    document = Document()
    document.add_paragraph("본문 유지")
    document.sections[0].header.add_paragraph("머리말")
    document.sections[0].footer.add_paragraph("바닥글")
    document.save(path)
    _replace_zip_members(path, {"word/header1.xml": b"<w:hdr"})

    # When the document is parsed.
    extraction, units = _parse(path)

    # Then the extraction is explicitly partial, the malformed part is
    # reported as a typed warning, and every other part remains intact.
    assert extraction.status == "partial"
    assert any(
        warning.startswith("PARTIAL_PARSE word/header1.xml:") for warning in extraction.warnings
    )
    assert any(unit.body == "본문 유지" for unit in units if unit.unit_type == "docx_paragraph")
    assert any(
        unit.body == "바닥글"
        for unit in units
        if unit.unit_type == "docx_header_footer" and unit.locator.data["part"] == "word/footer1.xml"
    )
    assert not any(
        unit.unit_type == "docx_header_footer" and unit.locator.data["part"] == "word/header1.xml"
        for unit in units
    )


def test_docx_parser_chunks_paragraph_locator_ranges_contiguously(tmp_path: Path) -> None:
    # Given a document with more paragraph text than fits in a single chunk.
    path = tmp_path / "long.docx"
    document = Document()
    for index in range(12):
        document.add_paragraph(f"문단 {index} " + "가" * 80)
    document.save(path)

    # When the document is parsed with a small per-unit character budget.
    _extraction, units = DocxParser(max_chars_per_unit=200).parse(
        path,
        artifact_id="art_long",
        document_id="doc_long",
        acl_scopes=["workspace:default"],
    )

    # Then multiple docx_paragraph units are emitted whose locator ranges are
    # contiguous and together cover every source paragraph exactly once.
    paragraph_units = [unit for unit in units if unit.unit_type == "docx_paragraph"]
    assert len(paragraph_units) > 1
    covered: list[int] = []
    for unit in paragraph_units:
        start = unit.locator.data["start_paragraph"]
        end = unit.locator.data["end_paragraph"]
        assert start <= end
        covered.extend(range(start, end + 1))
    assert covered == list(range(1, 13))


def test_docx_parser_returns_typed_error_for_corrupt_zip(tmp_path: Path) -> None:
    # Given bytes that are not a valid ZIP archive at all.
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a zip archive")

    # When the parser attempts to open the archive.
    # Then callers receive the stable parser error instead of a raw zip exception.
    with pytest.raises(ParserError, match="DOCX parse failed"):
        _parse(path)


def test_docx_parser_returns_typed_error_for_malformed_primary_document(tmp_path: Path) -> None:
    # Given a well-formed ZIP whose primary document part is malformed XML.
    path = tmp_path / "malformed.docx"
    document = Document()
    document.add_paragraph("본문")
    document.save(path)
    _replace_zip_members(path, {"word/document.xml": b"<w:document"})

    # When the parser reaches the malformed primary part.
    # Then callers receive the stable parser error instead of an XML implementation error.
    with pytest.raises(ParserError, match="DOCX parse failed"):
        _parse(path)


def test_docx_parser_quality_degrades_on_replacement_character_content(tmp_path: Path) -> None:
    # Given a clean baseline document and an otherwise identical document
    # whose text was corrupted into Unicode replacement characters.
    clean_path = tmp_path / "clean.docx"
    clean_document = Document()
    clean_document.add_paragraph("정상적인 본문 내용입니다")
    clean_document.save(clean_path)

    corrupt_path = tmp_path / "corrupt_text.docx"
    corrupt_document = Document()
    corrupt_document.add_paragraph("����������")
    corrupt_document.save(corrupt_path)

    # When both documents are parsed.
    clean_extraction, _clean_units = _parse(clean_path)
    corrupt_extraction, _corrupt_units = _parse(corrupt_path)

    # Then replacement-character content is scored lower than clean content.
    assert clean_extraction.quality_score is not None
    assert corrupt_extraction.quality_score is not None
    assert corrupt_extraction.quality_score < clean_extraction.quality_score


def test_docx_parser_quality_is_not_tanked_by_normal_blank_paragraph_spacing(
    tmp_path: Path,
) -> None:
    # Given a document with one real paragraph and fifty blank ones (normal
    # spacing/formatting, not an extraction failure).
    path = tmp_path / "blank_spacing.docx"
    document = Document()
    document.add_paragraph("정상적인 본문 내용입니다")
    for _ in range(50):
        document.add_paragraph("")
    document.save(path)

    # When the document is parsed.
    extraction, _units = _parse(path)

    # Then quality stays near the clean-document ceiling instead of being
    # tanked by paragraph density (1 real / 51 total previously scored
    # ~0.018).
    assert extraction.quality_score >= 0.9
    # Given a DOCX whose primary document part is a small-on-disk, highly
    # compressible payload that would expand to gigabytes in memory (the
    # same style of attack XlsxShallowParser and PptxParser already guard
    # against via a zip-bomb ratio/size check).
    path = tmp_path / "bomb.docx"
    document_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b"<w:body><w:p><w:r><w:t>" + b"A" * (50 * 1024 * 1024) + b"</w:t></w:r></w:p>"
        b"</w:body></w:document>"
    )
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        archive.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
            'Target="word/document.xml"/></Relationships>',
        )
        archive.writestr("word/document.xml", document_xml)

    # When the parser inspects the package.
    # Then it stops before expanding the archive into memory, mirroring the
    # XLSX/PPTX zip-bomb guard instead of decompressing an unbounded payload.
    with pytest.raises(ValidationError, match="decompression limits"):
        _parse(path)


def test_docx_parser_rejects_excessively_nested_document_xml(tmp_path: Path) -> None:
    # Given a small, low-compression-ratio DOCX whose primary document part
    # wraps a single run in thousands of nested elements. ET.fromstring
    # parses this fine (expat is iterative), but the hand-written run/textbox
    # walkers recurse in Python and would otherwise crash with an uncaught
    # RecursionError instead of a typed parser error.
    path = tmp_path / "deepnest.docx"
    depth = 3000
    open_tags = "".join(
        f'<w:x xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="v{i}">'
        for i in range(depth)
    )
    close_tags = "</w:x>" * depth
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body><w:p>" + open_tags + "<w:r><w:t>hi</w:t></w:r>" + close_tags + "</w:p></w:body>"
        "</w:document>"
    ).encode("utf-8")
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        archive.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
            'Target="word/document.xml"/></Relationships>',
        )
        archive.writestr("word/document.xml", document_xml)

    # When the parser walks the deeply nested run.
    # Then it raises a typed parser error instead of an uncaught RecursionError.
    with pytest.raises(ParserError, match="nesting exceeds"):
        _parse(path)


def _write_minimal_docx(path: Path, document_xml: bytes) -> None:
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        archive.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
            'Target="word/document.xml"/></Relationships>',
        )
        archive.writestr("word/document.xml", document_xml)


def test_docx_parser_recovers_table_nested_inside_a_cell_inside_a_table(tmp_path: Path) -> None:
    # Given a table whose single cell contains another, fully nested table
    # (table-in-cell-in-table) instead of only a paragraph.
    path = tmp_path / "nested_table.docx"
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        "<w:tbl><w:tr><w:tc>"
        "<w:p><w:r><w:t>바깥 텍스트</w:t></w:r></w:p>"
        "<w:tbl><w:tr><w:tc>"
        "<w:p><w:r><w:t>안쪽 값</w:t></w:r></w:p>"
        "</w:tc></w:tr></w:tbl>"
        "</w:tc></w:tr></w:tbl>"
        "</w:body></w:document>"
    ).encode()
    _write_minimal_docx(path, document_xml)

    # When the document is parsed.
    extraction, units = _parse(path)

    # Then a single top-level table unit is emitted (the nested table is not
    # a direct child of the body) whose rendered text recovers both the
    # outer and the inner table's content, and the nesting is counted.
    table_units = [unit for unit in units if unit.unit_type == "docx_table"]
    assert len(table_units) == 1
    assert "바깥 텍스트" in table_units[0].body
    assert "안쪽 값" in table_units[0].body
    assert table_units[0].metadata["nested_table_count"] == 1
    assert extraction.metadata["table_count"] == 1
    assert extraction.metadata["nested_table_count"] == 1


def test_docx_parser_recovers_textbox_nested_inside_a_table_inside_a_textbox(
    tmp_path: Path,
) -> None:
    # Given a text box whose content includes a table, one of whose cells
    # holds another, independently nested text box (textbox -> table ->
    # textbox), built with plain DrawingML (no mc:AlternateContent wrapper)
    # so the test isolates the txbxContent-nesting bug from Choice/Fallback
    # de-duplication.
    path = tmp_path / "nested_textbox.docx"
    document = Document()
    document.add_paragraph("OUTERPLACEHOLDER")
    document.save(path)
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")

    inner_text = "안쪽 텍스트상자"
    outer_text = "바깥 텍스트상자"
    inner_drawing = (
        "<w:p><w:r>"
        '<w:drawing><wp:anchor><wp:docPr id="2" name="TextBox 2"/>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        "<wps:wsp><wps:txbx><w:txbxContent>"
        f"<w:p><w:r><w:t>{inner_text}</w:t></w:r></w:p>"
        "</w:txbxContent></wps:txbx></wps:wsp>"
        "</a:graphicData></a:graphic></wp:anchor></w:drawing>"
        "</w:r></w:p>"
    )
    outer_drawing = (
        "<w:p><w:r>"
        '<w:drawing><wp:anchor><wp:docPr id="1" name="TextBox 1"/>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        "<wps:wsp><wps:txbx><w:txbxContent>"
        f"<w:p><w:r><w:t>{outer_text}</w:t></w:r></w:p>"
        "<w:tbl><w:tr><w:tc>" + inner_drawing + "</w:tc></w:tr></w:tbl>"
        "</w:txbxContent></wps:txbx></wps:wsp>"
        "</a:graphicData></a:graphic></wp:anchor></w:drawing>"
        "</w:r></w:p>"
    ).encode()
    document_xml = document_xml.replace(
        b"<w:p><w:r><w:t>OUTERPLACEHOLDER</w:t></w:r></w:p>", outer_drawing
    )
    _replace_zip_members(path, {"word/document.xml": document_xml})

    # When the document is parsed.
    extraction, units = _parse(path)

    # Then both the outer and the inner text box are recovered as their own
    # units (previously the walk returned immediately on finding a
    # txbxContent, so the inner one was silently dropped), the inner text is
    # not duplicated into the outer unit's body, and the nesting is counted.
    textbox_units = [unit for unit in units if unit.unit_type == "docx_textbox"]
    assert len(textbox_units) == 2
    assert textbox_units[0].body == outer_text
    assert textbox_units[1].body == inner_text
    assert inner_text not in textbox_units[0].body
    assert extraction.metadata["textbox_count"] == 2
    assert extraction.metadata["nested_textbox_count"] == 1


def test_docx_parser_captures_body_level_alternate_content_paragraph(tmp_path: Path) -> None:
    # Given a body-level paragraph wrapped in mc:AlternateContent (Word
    # sometimes wraps content that differs between application versions;
    # only text boxes handled this construct before this fix).
    path = tmp_path / "alternate_paragraph.docx"
    document = Document()
    document.add_paragraph("PLACEHOLDER")
    document.save(path)
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")
    wrapped_paragraph = (
        "<mc:AlternateContent>"
        '<mc:Choice Requires="w16se">'
        "<w:p><w:r><w:t>대체 콘텐츠 본문</w:t></w:r></w:p>"
        "</mc:Choice>"
        "<mc:Fallback>"
        "<w:p><w:r><w:t>대체 콘텐츠 폴백</w:t></w:r></w:p>"
        "</mc:Fallback>"
        "</mc:AlternateContent>"
    ).encode()
    document_xml = document_xml.replace(
        b"<w:p><w:r><w:t>PLACEHOLDER</w:t></w:r></w:p>", wrapped_paragraph
    )
    _replace_zip_members(path, {"word/document.xml": document_xml})

    # When the document is parsed.
    _extraction, units = _parse(path)

    # Then the Choice paragraph's text is captured as ordinary body text
    # instead of being silently dropped because it is not a direct "p"
    # child of the body.
    assert any("대체 콘텐츠 본문" in unit.body for unit in units)
    assert not any("대체 콘텐츠 폴백" in unit.body for unit in units)


def test_docx_heading_style_level_is_capped_to_nine() -> None:
    # Given a bogus two-digit heading style id (Word only defines levels 1-9).
    # When the style id is resolved to a heading level.
    # Then it is rejected instead of yielding a nonsensical level 99.
    assert docx_module._heading_level("Heading99") is None
    assert docx_module._heading_level("Heading9") == 9
    assert docx_module._heading_level("Heading1") == 1


def test_registry_selects_docx_parser_for_docx_extension(tmp_path: Path) -> None:
    # Given the shared parser registry configured with defaults.
    from kip.adapters.parsers.registry import ParserRegistry
    from kip.settings import Settings

    path = tmp_path / "status.docx"
    document = Document()
    document.add_paragraph("본문")
    document.save(path)
    settings = Settings(project_root=tmp_path, config_path=tmp_path / "kip.toml", raw={})

    # When the registry resolves a .docx artifact.
    parser = ParserRegistry.from_settings(settings).find(path)

    # Then the structural DOCX adapter owns the artifact.
    assert parser.name == "docx-xml"
